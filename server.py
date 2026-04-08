import os
import datetime
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import threading

app = Flask(__name__, static_folder='.')
CORS(app)

# --- CONFIGURATION ---
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
CLASSEUR_PATH = os.path.abspath(os.path.join(BASE_DIR, '..'))
DATABASE_PATH = os.path.join(BASE_DIR, 'caisse.db')

app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{DATABASE_PATH}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# --- MODELS ---
class Transaction(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    date = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    type = db.Column(db.String(10), nullable=False) # 'income' or 'expense'
    amount = db.Column(db.Float, nullable=False)
    description = db.Column(db.String(200), nullable=False)
    category = db.Column(db.String(50), nullable=False)
    reference = db.Column(db.String(50), nullable=False)
    initiator = db.Column(db.String(100), nullable=False)
    balance_after = db.Column(db.Float)

class ReportLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    report_type = db.Column(db.String(20)) # 'monthly', 'quarterly'
    period = db.Column(db.String(50)) # '2026-01', '2026-Q1'
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)

with app.app_context():
    db.create_all()

# --- HELPERS ---
MONTH_NAMES = {
    1: "Janvier", 2: "Février", 3: "Mars", 4: "Avril", 5: "Mai", 6: "Juin",
    7: "Juillet", 8: "Août", 9: "Septembre", 10: "Octobre", 11: "Novembre", 12: "Décembre"
}

def get_report_dir(year, month_name=None):
    path = os.path.join(CLASSEUR_PATH, "Rapports", str(year))
    if month_name:
        path = os.path.join(path, month_name)
    os.makedirs(path, exist_ok=True)
    return path

def format_curr(amount):
    return f"{amount:,.0f} F CFA".replace(",", " ")

def generate_transaction_docx(t):
    doc = Document()
    year = t.date.year
    month_name = MONTH_NAMES[t.date.month]
    report_dir = get_report_dir(year, month_name)
    
    # Professional Header
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = "ASSOCIATION - RECU D'OPÉRATION"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(f"Détails de l'Opération: {t.reference}", 0)
    
    table = doc.add_table(rows=0, cols=2)
    rows = [
        ("Date & Heure", t.date.strftime("%d/%m/%Y %H:%M")),
        ("Référence", t.reference),
        ("Initiateur", t.initiator),
        ("Type", "ENTRÉE" if t.type == 'income' else "SORTIE"),
        ("Catégorie", t.category),
        ("Libellé", t.description),
        ("Montant", format_curr(t.amount)),
        ("Solde après", format_curr(t.balance_after or 0))
    ]
    for label, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value)

    doc.add_paragraph("\n\nSignature du Responsable:")
    
    filename = f"transaction_{t.reference}_{t.id}.docx"
    doc.save(os.path.join(report_dir, filename))
    return filename

def generate_periodic_report(month=None, year=None, quarter=None):
    current_year = year or datetime.datetime.now().year
    doc = Document()
    
    if quarter:
        title = f"BILAN TRIMESTRIEL - Q{quarter} {current_year}"
        months_in_q = range((quarter-1)*3 + 1, quarter*3 + 1)
        transactions = Transaction.query.filter(
            db.extract('year', Transaction.date) == current_year,
            db.extract('month', Transaction.date).in_(months_in_q)
        ).all()
        report_dir = os.path.join(CLASSEUR_PATH, "Rapports", str(current_year))
        filename = f"bilan_trimestriel_Q{quarter}_{current_year}.docx"
    else:
        month_name = MONTH_NAMES[month]
        title = f"RAPPORT MENSUEL - {month_name.upper()} {current_year}"
        transactions = Transaction.query.filter(
            db.extract('year', Transaction.date) == current_year,
            db.extract('month', Transaction.date) == month
        ).all()
        report_dir = get_report_dir(current_year, month_name)
        filename = f"rapport_mensuel_{month_name}_{current_year}.docx"

    doc.add_heading(title, 0)
    
    sum_in = sum(t.amount for t in transactions if t.type == 'income')
    sum_out = sum(t.amount for t in transactions if t.type == 'expense')
    
    doc.add_heading("Résumé Financier", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Total des Entrées: ").bold = True
    p.add_run(format_curr(sum_in))
    p = doc.add_paragraph()
    p.add_run(f"Total des Sorties: ").bold = True
    p.add_run(format_curr(sum_out))
    p = doc.add_paragraph()
    p.add_run(f"Solde Net Période: ").bold = True
    p.add_run(format_curr(sum_in - sum_out))

    doc.add_heading("Détail des Transactions", level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Réf'
    hdr_cells[2].text = 'Libellé'
    hdr_cells[3].text = 'Entrée'
    hdr_cells[4].text = 'Sortie'

    for t in sorted(transactions, key=lambda x: x.date):
        row_cells = table.add_row().cells
        row_cells[0].text = t.date.strftime("%d/%m/%Y")
        row_cells[1].text = t.reference
        row_cells[2].text = t.description
        row_cells[3].text = format_curr(t.amount) if t.type == 'income' else "-"
        row_cells[4].text = format_curr(t.amount) if t.type == 'expense' else "-"

    os.makedirs(report_dir, exist_ok=True)
    doc.save(os.path.join(report_dir, filename))
    return filename

def check_and_generate_automated_reports():
    with app.app_context():
        now = datetime.datetime.now()
        # Monthly check for previous month
        prev_month_date = now.replace(day=1) - datetime.timedelta(days=1)
        m, y = prev_month_date.month, prev_month_date.year
        period_str = f"{y}-{m:02d}"
        
        if not ReportLog.query.filter_by(report_type='monthly', period=period_str).first():
            print(f"Generating automated monthly report for {period_str}")
            generate_periodic_report(month=m, year=y)
            db.session.add(ReportLog(report_type='monthly', period=period_str))
            db.session.commit()

        # Quarterly check
        q = (now.month - 1) // 3
        if q > 0: # If we are in Q2/Q3/Q4, check prev quarter
            prev_q = q
            period_q = f"{now.year}-Q{prev_q}"
            if not ReportLog.query.filter_by(report_type='quarterly', period=period_q).first():
                generate_periodic_report(year=now.year, quarter=prev_q)
                db.session.add(ReportLog(report_type='quarterly', period=period_q))
                db.session.commit()

# --- API ROUTES ---
@app.route('/')
def index():
    return send_from_directory('.', 'index.html')

@app.route('/<path:path>')
def static_files(path):
    return send_from_directory('.', path)

@app.route('/api/transactions', methods=['GET'])
def get_transactions():
    query = Transaction.query
    # Filters could be added here (month, type, etc)
    transactions = query.order_by(Transaction.date.desc()).all()
    return jsonify([{
        'id': t.id, 'date': t.date.strftime("%Y-%m-%d"), 'type': t.type,
        'amount': t.amount, 'description': t.description, 'category': t.category,
        'reference': t.reference, 'initiator': t.initiator, 'balance_after': t.balance_after
    } for t in transactions])

@app.route('/api/transactions', methods=['POST'])
def add_transaction():
    data = request.json
    last_t = Transaction.query.order_by(Transaction.id.desc()).first()
    prev_balance = last_t.balance_after if last_t else 0
    
    amount = float(data['amount'])
    new_balance = prev_balance + (amount if data['type'] == 'income' else -amount)
    
    new_t = Transaction(
        date=datetime.datetime.strptime(data['date'], "%Y-%m-%d") if data.get('date') else datetime.datetime.now(),
        type=data['type'],
        amount=amount,
        description=data['description'],
        category=data['category'],
        reference=data['reference'],
        initiator=data['initiator'],
        balance_after=new_balance
    )
    db.session.add(new_t)
    db.session.commit()
    
    # Generate automatic docx
    generate_transaction_docx(new_t)
    
    return jsonify({'success': True, 'id': new_t.id})

@app.route('/api/stats')
def get_stats():
    last_t = Transaction.query.order_by(Transaction.id.desc()).first()
    balance = last_t.balance_after if last_t else 0
    total_in = db.session.query(db.func.sum(Transaction.amount)).filter(Transaction.type == 'income').scalar() or 0
    total_out = db.session.query(db.func.sum(Transaction.amount)).filter(Transaction.type == 'expense').scalar() or 0
    return jsonify({
        'balance': balance,
        'total_income': total_in,
        'total_expense': total_out
    })

@app.route('/api/export/excel', methods=['POST'])
def export_excel():
    transactions = Transaction.query.all()
    df = pd.DataFrame([{
        'Date': t.date, 'Type': t.type, 'Montant': t.amount,
        'Description': t.description, 'Catégorie': t.category,
        'Référence': t.reference, 'Initiateur': t.initiator, 'Solde': t.balance_after
    } for t in transactions])
    
    path = os.path.join(BASE_DIR, "export_caisse.xlsx")
    df.to_excel(path, index=False)
    return send_file(path, as_attachment=True)

@app.route('/api/reports')
def list_reports():
    reports = []
    # Simplified: walk Rapports dir
    root = os.path.join(CLASSEUR_PATH, "Rapports")
    if os.path.exists(root):
        for r, d, f in os.walk(root):
            for file in f:
                if file.endswith(".docx"):
                    reports.append({
                        'name': file,
                        'path': os.path.relpath(os.path.join(r, file), root)
                    })
    return jsonify(reports[:50]) # Limit to 50

if __name__ == '__main__':
    # Run automated check on start
    threading.Thread(target=check_and_generate_automated_reports).start()
    app.run(port=5000, debug=True)
