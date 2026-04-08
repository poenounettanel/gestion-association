import os
import sys
import threading
import time
import socket
import subprocess
import webbrowser
import datetime
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

# Handle bundled asset paths
def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

app = Flask(__name__, static_folder=resource_path('.'))
CORS(app)

# --- CONFIGURATION (Paths absolute for portability) ---
# We use the location of the EXE or script as base
EXECUTABLE_DIR = os.path.dirname(os.path.abspath(sys.executable if hasattr(sys, 'frozen') else __file__))
CLASSEUR_ROOT = EXECUTABLE_DIR # The app will create folders next to it

# Database stays next to the app
DATABASE_PATH = os.path.join(EXECUTABLE_DIR, 'caisse_pro.db')
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{DATABASE_PATH}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# --- MODELS ---
class Transaction(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    date = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    type = db.Column(db.String(10), nullable=False)
    amount = db.Column(db.Float, nullable=False)
    description = db.Column(db.String(200), nullable=False)
    category = db.Column(db.String(50), nullable=False)
    reference = db.Column(db.String(50), nullable=False)
    initiator = db.Column(db.String(100), nullable=False)
    balance_after = db.Column(db.Float)

class ReportLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    report_type = db.Column(db.String(20))
    period = db.Column(db.String(50))
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)

class Category(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), unique=True, nullable=False)

with app.app_context():
    db.create_all()
    if Category.query.count() == 0:
        default_cats = ["Cotisations", "Dons", "Quêtes", "Ventes exceptionnelles", "Achats courants", "Secours/Aides", "Transports", "Frais de fonctionnement"]
        for cat in default_cats:
            db.session.add(Category(name=cat))
        db.session.commit()

# --- HELPERS (Copied from server.py) ---
MONTH_NAMES = {1: "Janvier", 2: "Février", 3: "Mars", 4: "Avril", 5: "Mai", 6: "Juin",
               7: "Juillet", 8: "Août", 9: "Septembre", 10: "Octobre", 11: "Novembre", 12: "Décembre"}

def get_report_dir(year, month_name=None):
    path = os.path.join(CLASSEUR_ROOT, "Rapports", str(year))
    if month_name:
        path = os.path.join(path, month_name)
    os.makedirs(path, exist_ok=True)
    return path

def format_curr(amount):
    return f"{amount:,.0f} F CFA".replace(",", " ")

def generate_transaction_docx(t):
    doc = Document()
    year = t.date.year
    month_name = MONTH_NAMES[t.date.month]
    report_dir = get_report_dir(year, month_name)
    doc.add_heading(f"RECU D'OPÉRATION: {t.reference}", 0)
    table = doc.add_table(rows=0, cols=2)
    rows = [("Date", t.date.strftime("%d/%m/%Y")), ("Référence", t.reference), ("Initiateur", t.initiator),
            ("Type", t.type), ("Catégorie", t.category), ("Libellé", t.description), ("Montant", format_curr(t.amount))]
    for l, v in rows:
        r = table.add_row().cells
        r[0].text, r[1].text = l, str(v)
    doc.save(os.path.join(report_dir, f"recu_{t.reference}.docx"))

# --- ROUTES ---
@app.route('/')
def index(): return send_from_directory(resource_path('.'), 'index.html')
@app.route('/<path:path>')
def static_files(path): return send_from_directory(resource_path('.'), path)

@app.route('/api/transactions', methods=['GET', 'POST'])
def transactions_api():
    if request.method == 'POST':
        data = request.json
        last_t = Transaction.query.order_by(Transaction.id.desc()).first()
        prev = last_t.balance_after if last_t else 0
        amount = float(data['amount'])
        new_bal = prev + (amount if data['type'] == 'income' else -amount)
        cat_name = data['category'].strip()
        
        # Add category dynamically if it does not exist
        if not Category.query.filter_by(name=cat_name).first():
            db.session.add(Category(name=cat_name))
            db.session.commit()

        new_t = Transaction(date=datetime.datetime.strptime(data['date'], "%Y-%m-%d"), type=data['type'], amount=amount,
                            description=data['description'], category=cat_name, reference=data['reference'],
                            initiator=data['initiator'], balance_after=new_bal)
        db.session.add(new_t)
        db.session.commit()
        generate_transaction_docx(new_t)
        return jsonify({'success': True})
    ts = Transaction.query.order_by(Transaction.date.desc()).all()
    return jsonify([{'id':t.id,'date':t.date.strftime("%Y-%m-%d"),'type':t.type,'amount':t.amount,'description':t.description,
                     'category':t.category,'reference':t.reference,'initiator':t.initiator,'balance_after':t.balance_after} for t in ts])

@app.route('/api/categories', methods=['GET'])
def get_categories():
    cats = Category.query.order_by(Category.name).all()
    return jsonify([c.name for c in cats])

@app.route('/api/stats')
def stats():
    last_t = Transaction.query.order_by(Transaction.id.desc()).first()
    return jsonify({'balance': last_t.balance_after if last_t else 0,
                    'total_income': db.session.query(db.func.sum(Transaction.amount)).filter(Transaction.type == 'income').scalar() or 0,
                    'total_expense': db.session.query(db.func.sum(Transaction.amount)).filter(Transaction.type == 'expense').scalar() or 0})

@app.route('/api/reports')
def list_reports():
    reports = []
    root = os.path.join(CLASSEUR_ROOT, "Rapports")
    if os.path.exists(root):
        for r, d, f in os.walk(root):
            for file in f:
                if file.endswith(".docx"):
                    reports.append({'name': file, 'path': os.path.relpath(os.path.join(r, file), root)})
    return jsonify(reports)

@app.route('/api/export/excel', methods=['POST'])
def export():
    ts = Transaction.query.all()
    df = pd.DataFrame([{'Date':t.date,'Type':t.type,'Montant':t.amount,'Description':t.description,'Catégorie':t.category,'Référence':t.reference,'Initiateur':t.initiator,'Solde':t.balance_after} for t in ts])
    path = os.path.join(EXECUTABLE_DIR, "export_caisse.xlsx")
    df.to_excel(path, index=False)
    return send_file(path, as_attachment=True)

# --- LAUNCHER LOGIC ---
def is_port_in_use(port):
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        return s.connect_ex(('localhost', port)) == 0

def run_flask():
    app.run(port=5000, threaded=True)

if __name__ == '__main__':
    if not is_port_in_use(5000):
        threading.Thread(target=run_flask, daemon=True).start()
        time.sleep(2)
    
    url = "http://localhost:5000"
    edge_path = r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
    if os.path.exists(edge_path):
        subprocess.Popen([edge_path, f"--app={url}"])
    else:
        webbrowser.open(url)
    
    # Keep main alive while server runs
    while True: time.sleep(10)
